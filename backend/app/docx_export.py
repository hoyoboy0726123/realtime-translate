"""Render a meeting summary (markdown) as a Word .docx report.

The summary the LLM produces is markdown — headings, bullet / numbered lists,
**bold** spans and the occasional table. This converts that subset into a
properly formatted Word document.
"""
from __future__ import annotations

import re
from io import BytesIO

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBER = re.compile(r"^\s*\d+\.\s+(.*)$")


def _add_runs(paragraph, text: str) -> None:
    """Append text to a paragraph, rendering **bold** spans as bold runs."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.count("|") >= 2


def _is_table_separator(line: str) -> bool:
    return "-" in line and bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line))


def _cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def summary_to_docx(title: str, summary_md: str) -> bytes:
    """Build a .docx report from a session title and a markdown summary."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, 0)

    lines = summary_md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # Table: a header row, a separator row, then data rows.
        if (_is_table_row(line) and i + 1 < len(lines)
                and _is_table_separator(lines[i + 1])):
            header = _cells(line)
            data = []
            i += 2
            while i < len(lines) and _is_table_row(lines[i]):
                data.append(_cells(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for c, text in enumerate(header):
                table.rows[0].cells[c].text = text
            for row in data:
                cells = table.add_row().cells
                for c in range(len(header)):
                    cells[c].text = row[c] if c < len(row) else ""
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif (m := _BULLET.match(line)) is not None:
            _add_runs(doc.add_paragraph(style="List Bullet"), m.group(1))
        elif (m := _NUMBER.match(line)) is not None:
            _add_runs(doc.add_paragraph(style="List Number"), m.group(1))
        else:
            _add_runs(doc.add_paragraph(), stripped)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
